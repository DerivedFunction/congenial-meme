import docx
import json
import os
from io import BytesIO
from docx.shared import Pt
from typing import Dict, Union, List, Tuple

def edit_word_tables(input_docx_path: str, json_data: str) -> Tuple[bytes, Dict[str, Union[str, List]]]:
    """
    Finds tables in a Word document, identifies cells with 'EDIT_' or 'EDIT' markers, and updates them with JSON values,
    setting font to Times New Roman, 9pt. Returns the modified document as bytes.
    
    Args:
        input_docx_path (str): Path to the input Word (.docx) file.
        json_data (str): JSON string with keys matching 'EDIT_' placeholders (e.g., 'EDIT_DOR').
            Example:
            {
                "EDIT_lastName": "LI",
                "EDIT_FirstName": "DENNY",
                "Edit_MI": "",
                ...
                "EDIT_MOSDESC": "...",
                "EDIT_Topics": "asdf",
                ...
            }
    
    Returns:
        Tuple[bytes, Dict]: (Modified document as bytes, response metadata with status and messages).
    """
    response = {"status": "success", "messages": []}
    
    try:
        # Validate input Word document path
        if not os.path.exists(input_docx_path):
            response["status"] = "error"
            response["messages"].append(f"Word document not found: {input_docx_path} (current working directory: {os.getcwd()})")
            return b"", response
        
        # Parse JSON input
        try:

            field_values = json.loads(json_data)
            if not isinstance(field_values, dict):
                raise ValueError("JSON input must be a dictionary.")
        except json.JSONDecodeError as e:
            response["status"] = "error"
            response["messages"].append(f"Invalid JSON format: {str(e)} JSON: {json_data}")
            return b"", response
        
        # Load the Word document
        try:
            print(f"Reading Word document: {input_docx_path}")
            doc = docx.Document(input_docx_path)
        except Exception as e:
            response["status"] = "error"
            response["messages"].append(f"Failed to read Word document: {str(e)}")
            return b"", response
        
        # Lists to track tables and cells
        edit_cells = []  # Cells with "EDIT_" or "EDIT"
        updated_cells = []  # Cells updated with JSON values
        all_tables = []  # All table details
        
        # Iterate through all tables in the document
        for table_idx, table in enumerate(doc.tables):
            table_info = {
                "table_index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }
            print(f"\nTable {table_idx}: {table_info['rows']} rows, {table_info['columns']} columns")
            
            # Iterate through each cell in the table
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_info = {
                        "row": row_idx + 1,  # 1-based indexing to match provided output
                        "column": cell_idx + 1,
                        "text": cell_text
                    }
                    table_info["cells"].append(cell_info)
                    
                    # Check for "EDIT_" or "EDIT" marker (case-insensitive)
                    if cell_text.lower().startswith("edit_") or cell_text.lower() == "edit":
                        edit_cells.append(f"Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}): {cell_text}")
                        print(f"  Found editable cell in Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}): {cell_text}")
                        
                        # Update cell if it matches a JSON key
                        if cell_text in field_values:
                            new_value = field_values[cell_text]
                            if new_value is None:
                                response["messages"].append(f"Skipped cell in Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}): null value provided for {cell_text}")
                                continue
                            try:
                                new_value_str = str(new_value)
                                # Clear existing content
                                for paragraph in cell.paragraphs:
                                    paragraph.clear()
                                # Add new paragraph with font settings
                                paragraph = cell.add_paragraph(new_value_str)
                                for run in paragraph.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(9)
                                updated_cells.append(f"Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}): {new_value_str}")
                                response["messages"].append(f"Updated cell in Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}) with value: {new_value_str} (font: Times New Roman, 9pt)")
                            except Exception as e:
                                response["messages"].append(f"Failed to update cell in Table {table_idx}, Cell ({row_idx + 1}, {cell_idx + 1}): {str(e)}")
                    else:
                        print(f"  Cell ({row_idx + 1}, {cell_idx + 1}): {cell_text}")
            
            all_tables.append(table_info)
        
        # Save the modified document to a BytesIO stream
        output_stream = BytesIO()
        try:
            doc.save(output_stream)
            doc_bytes = output_stream.getvalue()
            output_stream.close()
            response["messages"].append("Generated modified Word document as byte stream")
        except Exception as e:
            response["status"] = "error"
            response["messages"].append(f"Failed to generate document bytes: {str(e)}")
            return b"", response
        
        # Add table and cell information to response
        response["messages"].append(f"Found {len(doc.tables)} tables in the document.")
        if edit_cells:
            response["messages"].append(f"Editable cells: {', '.join(edit_cells)}")
        else:
            response["messages"].append("No cells with 'EDIT_' or 'EDIT' found.")
        if updated_cells:
            response["messages"].append(f"Updated cells: {', '.join(updated_cells)}")
        else:
            response["messages"].append("No cells updated.")
        response["messages"].append("Table details:")
        for table_info in all_tables:
            response["messages"].append(f"Table {table_info['table_index']}: {table_info['rows']} rows, {table_info['columns']} columns")
            for cell_info in table_info["cells"]:
                response["messages"].append(f"  Cell ({cell_info['row']}, {cell_info['column']}): {cell_info['text']}")
        
        print("\nResponse messages:")
        for msg in response["messages"]:
            print(msg)
        
        return doc_bytes, response
    
    except Exception as e:
        response["status"] = "error"
        response["messages"].append(f"Unexpected error: {str(e)}")
        print("Error:", response["messages"])
        return b"", response